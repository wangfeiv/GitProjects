#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

OUT = '/home/wangfei/.openclaw/workspace/多多四年级数学错题专项复习卷-提高版.docx'


def set_font(run, size=12, bold=False):
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(size)
    run.bold = bold


def add_p(doc, text='', size=12, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def add_lines(doc, lines, size=12):
    for line in lines:
        add_p(doc, line, size=size)


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)

    add_p(doc, '多多四年级数学错题专项复习卷（提高版）', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, '姓名：__________    日期：__________    用时：__________    得分：__________', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, '说明：本卷是提高版，题目更综合，要求更会辨析、更会变式。全部为纯文本，可直接打印。')

    add_p(doc, '一、填空题（每空 4 分，共 32 分）', bold=True)
    add_lines(doc, [
        '1. 用 0、3、5、9 四个数字和小数点组成数，每个数字都要用一次。',
        '   最大的一位小数是____________，最小的两位小数是____________。',
        '2. 一个三位小数四舍五入后是 9.70，这个三位小数最大是____________，最小是____________。',
        '3. 一辆汽车每小时行 a 千米，行了 b 小时，路程是____________千米。',
        '   如果路程用 s 表示，那么 s=____________。',
        '4. 一个长方体，长 14 厘米，宽 11 厘米，高 5 厘米。',
        '   最大一个面的面积是____________平方厘米，最小一个面的面积是____________平方厘米。',
        '5. 苹果每千克 x 元，梨每千克 y 元，买 5 千克苹果比买 3 千克梨便宜____________元。',
        '6. 小华每分钟走 t 米，12 分钟走____________米；一共走 360 米，需要____________分钟。',
    ])

    add_p(doc, '二、判断与选择（每题 4 分，共 24 分）', bold=True)
    add_lines(doc, [
        '7. 判断：25×(40×4) 可以写成 25×40+25×4。 （    ）',
        '8. 判断：356-56-44 可以写成 356-(56+44)。 （    ）',
        '9. 甲数是 m，比乙数的 8 倍多 n，乙数是（    ）。',
        '   A. (m+n)÷8    B. (m-n)÷8    C. 8m-n    D. 8m+n',
        '10. 小明跑步用了 x 秒，小强比小明慢 4 秒，小强用了（    ）秒。',
        '   A. x-4    B. x+4    C. 4x    D. x÷4',
        '11. 聪聪把 (6+□)×18 错算成 6+□×18，错误结果与正确结果相差（    ）。',
        '   A. 96    B. 102    C. 108    D. 18',
        '12. 下面最适合用乘法结合律简便计算的是（    ）。',
        '   A. (120+8)×5    B. 125×16×25    C. 99×36    D. (80+4)×25',
    ])

    add_p(doc, '三、用简便方法计算（每题 6 分，共 36 分）', bold=True)
    add_lines(doc, [
        '13. 104×27',
        '14. 99×48',
        '15. 125×24',
        '16. 125×16×25',
        '17. (135+24)×4',
        '18. 478-78-22',
    ])

    add_p(doc, '四、解决问题（每题 4 分，共 8 分）', bold=True)
    add_lines(doc, [
        '19. 商店运进一批校服，进价每套 58 元，售价每套 76 元。全部卖完共赚了 540 元。',
        '    这批校服一共有多少套？',
        '20. 超市运来苹果和橙子各 12 箱。苹果每箱 28 千克，橙子每箱 32 千克。',
        '    一共运来多少千克水果？请用简便方法解答。',
    ])

    doc.add_page_break()

    add_p(doc, '参考答案与解析', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_p(doc, '一、填空题答案', bold=True)
    add_lines(doc, [
        '1. 最大的一位小数：953.0    最小的两位小数：30.59',
        '2. 最大：9.704    最小：9.695',
        '3. ab；ab',
        '4. 最大面：14×11=154    最小面：11×5=55',
        '5. 3y-5x',
        '6. 12t；360÷t',
    ])

    add_p(doc, '二、判断与选择答案', bold=True)
    add_lines(doc, [
        '7. 错。解析：括号里是乘法，应该用乘法结合律，不能写成分配律。',
        '8. 对。解析：连续减两个数，等于减去这两个数的和。',
        '9. B。因为 m=8×乙+n，所以 8×乙=m-n，乙=(m-n)÷8。',
        '10. B。慢表示用时更多，所以是 x+4。',
        '11. B。正确结果比错误结果多 6×18-6=108-6=102。',
        '12. B。125×16×25 最适合拆数凑整。',
    ])

    add_p(doc, '三、简便计算答案', bold=True)
    add_lines(doc, [
        '13. 104×27=(100+4)×27=100×27+4×27=2700+108=2808',
        '14. 99×48=(100-1)×48=100×48-1×48=4800-48=4752',
        '15. 125×24=125×(8×3)=(125×8)×3=1000×3=3000',
        '16. 125×16×25=125×(4×4)×25=(125×4)×(25×4)=500×100=50000',
        '17. (135+24)×4=135×4+24×4=540+96=636',
        '18. 478-78-22=478-(78+22)=478-100=378',
    ])

    add_p(doc, '四、解决问题答案', bold=True)
    add_lines(doc, [
        '19. 每套利润：76-58=18（元）',
        '    一共有：540÷18=30（套）',
        '    答：这批校服一共有 30 套。',
        '20. 方法一：28×12+32×12=(28+32)×12=60×12=720（千克）',
        '    答：一共运来 720 千克水果。',
    ])

    add_p(doc, '提高提醒：', bold=True)
    add_lines(doc, [
        '1. 看到“快”就想：时间更少；看到“慢”就想：时间更多。',
        '2. 看到“比……便宜多少”，一定是贵的减便宜的。',
        '3. 括号里如果是加减法，优先想分配律；括号里如果是乘法，优先想结合律。',
        '4. 添括号时要特别注意：括号前面是减号，括号里的符号要变。',
    ])

    add_p(doc, '出卷时间：' + datetime.now().strftime('%Y-%m-%d %H:%M'))
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build_doc()
