#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完全重新生成 四年级数学错题专项复习卷 纯文本版本
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def create_fresh_test():
    doc = Document()
    
    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(1.8)
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 标题
    title = doc.add_heading('四年级数学下册\n错题专项复习卷（第二套）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    
    # 信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('姓名：___________    用时：___________    得分：___________')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph('📋 本卷覆盖多多错题知识点：小数组合、四舍五入、用字母表示数、乘法分配律、乘法结合律、添括号变号、长方体面积。')
    doc.add_paragraph()
    
    # ========== 第一部分：填空题 ==========
    doc.add_heading('一、填空题（每空3分，共36分）', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('1. 用 3、0、6、8 四个数字和小数点组成（每个数字都要用，只能用一次）：')
    p = doc.add_paragraph()
    run = p.add_run('   - 最大的一位小数是 ____________')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小的两位小数是 ____________')
    
    p = doc.add_paragraph()
    run = p.add_run('2. 一个两位小数四舍五入后是 7.8，这个两位小数：')
    p = doc.add_paragraph()
    run = p.add_run('   - 最大是 ____________')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小是 ____________')
    
    p = doc.add_paragraph()
    run = p.add_run('3. 一个三位小数四舍五入后是 6.00，这个三位小数：')
    p = doc.add_paragraph()
    run = p.add_run('   - 最大是 ____________')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小是 ____________')
    
    p = doc.add_paragraph()
    run = p.add_run('4. 一块长方体木块，长 15 厘米，宽 10 厘米，高 6 厘米。')
    p = doc.add_paragraph()
    run = p.add_run('   - 最大一个面的面积是 ____________ 平方厘米')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小一个面的面积 is ____________ 平方厘米')
    
    p = doc.add_paragraph()
    run = p.add_run('5. 一辆汽车每小时行 v 千米，t 小时行 ____________ 千米。如果路程是 s，公式是 ____________。')
    
    p = doc.add_paragraph()
    run = p.add_run('6. 西瓜每千克 m 元，桃子每千克 n 元。买 6 千克西瓜比买 2 千克桃子便宜 ____________ 元。（用字母表示）')
    
    p = doc.add_paragraph()
    run = p.add_run('7. 小明每分钟走 a 米，15分钟走 ____________ 米，走 200米需要 ____________ 分钟。')
    
    doc.add_paragraph()
    
    # ========== 第二部分：选择题 ==========
    doc.add_heading('二、选择题（每题4分，共24分）', level=1)
    
    questions = [
        ('1. 甲数是 m，比乙数的5倍少n，乙数是（  ）', 
         ['A. 5m - n', 'B. (m + n) ÷ 5', 'C. (m - n) ÷ 5', 'D. m ÷ 5 - n'],
         ''),
        ('2. 50米短跑，小刚用了 x 秒，小明比小刚快 2 秒，小明用了（  ）秒',
         ['A. x + 2', 'B. x - 2', 'C. 2x', 'D. x ÷ 2'],
         ''),
        ('3. 正方形的周长是 C，它的边长是（  ）',
         ['A. 4C', 'B. C ÷ 4', 'C. C²', 'D. C - 4'],
         ''),
        ('4. 小明计算 (5 + □) × 24 时，错算成 5 + □ × 24，他算出的结果与正确结果相差（  ）',
         ['A. 100', 'B. 115', 'C. 120', 'D. 24'],
         ''),
        ('5. 学校买来足球和排球各 6 个，每个足球 75 元，每个排球 25 元，一共花了多少元？用简便方法计算，列式正确的是（  ）',
         ['A. 75 × 6 + 25 × 6', 'B. 6 × (75 + 25)', 'C. 75 × 25 × 6', 'D. 75 + 25 × 6'],
         ''),
        ('6. 与 102 × 35 计算结果相等的式子是（  ）',
         ['A. 100 × 35 + 2', 'B. 100 × 35 + 2 × 35', 'C. 100 + 2 × 35', 'D. (100 - 2) × 35'],
         ''),
    ]
    
    for num, question, options in questions:
        p = doc.add_paragraph()
        run = p.add_run(question)
        run.font.size = Pt(12)
        for opt in options:
            p = doc.add_paragraph()
            run = p.add_run(f'   {opt}')
            run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run('你的选择：____')
        run.font.size = Pt(11)
        doc.add_paragraph()
    
    # ========== 第三部分：简便计算 ==========
    doc.add_heading('三、用简便方法计算（每题5分，共30分）', level=1)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('💡 做题提示：')
    run.bold = True
    run = p.add_run('想一想，哪些题目用乘法分配律，哪些用乘法结合律？括号里面是加法还是乘法？')
    run.font.size = Pt(11)
    doc.add_paragraph()
    
    calc = [
        '1.  98 × 45              （提示：把98看成 100 - 2，用乘法分配律）',
        '2.  102 × 36             （提示：把102看成 100 + 2，用乘法分配律）',
        '3.  125 × 48             （提示：看到125，把48拆成 8 × 6，用乘法结合律）',
        '4.  25 × 125 × 32        （提示：看到125和25，把32拆成 4 × 8）',
        '5.  (30 + 4) × 25        （提示：乘法分配律，括号里每一项都要和25相乘）',
        '6.  456 - 56 - 44        （提示：连续减两个数，等于减这两个数的和）',
    ]
    for q in calc:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.size = Pt(12)
        p = doc.add_paragraph('\n')
    
    doc.add_paragraph()
    
    # ========== 第四部分：解决问题 ==========
    doc.add_heading('四、解决问题（10分）', level=1)
    p = doc.add_paragraph()
    run = p.add_run('服装店运进一批牛仔裤，进价每条 85 元，售价每条 135 元。全部卖完后总利润是 6000 元。这批牛仔裤一共有多少条？\n（提示：先算每条赚多少元，再算总共有多少条）')
    run.font.size = Pt(12)
    p = doc.add_paragraph('\n\n\n\n')
    
    # ========== 分页放答案 ==========
    doc.add_page_break()
    
    # 答案部分
    ans_title = doc.add_heading('参考答案与详细解析', 0)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # 填空题答案
    p = doc.add_heading('一、填空题', level=1)
    ans_fill = [
        ('1', '最大一位小数：863.0，最小两位小数：30.68', 
         '解析：一位小数，小数点后只有1位，剩下三个数字放整数部分，从大到小排列；'
         '两位小数，最小不能以0开头，所以3开头，从小到大排列。每个数字必须用完。'),
        ('2', '最大：7.84，最小：7.75',
         '解析：四舍五入规则：最大是"四舍"得到，所以第二位小数最大是4；最小是"五入"得到，第二位小数最小是5。记住：最大四舍，最小五入。'),
        ('3', '最大：6.004，最小：5.995',
         '解析：同理，6.005五入后变成6.01，所以最大只能是6.004。'),
        ('4', '最大面：15 × 10 = 150 → 150 平方厘米；最小面：6 × 10 = 60 → 60 平方厘米',
         '规律：最大面 = 最长两条边相乘，最小面 = 最短两条边相乘。'),
        ('5', '答案：vt，公式：s = vt',
         '公式：路程 = 速度 × 时间 → s = v × t = vt。'),
        ('6', '答案：2n - 6m',
         '"A比B便宜多少" = B总价 - A总价 → 2n - 6m。注意谁比谁，用大减小。'),
        ('7', '答案：15a 米，200 ÷ a 分钟',
         '路程 = 速度 × 时间；时间 = 路程 ÷ 速度。'),
    ]
    for num, ans, explain in ans_fill:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}. ')
        run.bold = True
        run = p.add_run(ans)
        p = doc.add_paragraph()
        run = p.add_run(f'   {explain}')
        run.font.size = Pt(11)
        doc.add_paragraph()
    
    doc.add_paragraph()
    
    # 选择题答案
    p = doc.add_heading('二、选择题', level=1)
    ans_choice = [
        ('1', 'B', '推导：甲数 = 乙数 × 5 - n → m = 5×乙 - n → 5×乙 = m + n → 乙 = (m + n) ÷ 5'),
        ('2', 'B', '"快" → 用时更少 → x - 2。快用减法，慢用加法。'),
        ('3', 'B', '正方形周长 = 边长 × 4 → 边长 = 周长 ÷ 4 = C ÷ 4'),
        ('4', 'B', '正确结果：(5+□)×24 = 5×24 + □×24 = 120 + □×24，'
         '错误结果：5 + □×24，相差：(120 + □×24) - (5 + □×24) = 115'),
        ('5', 'B', '(75 + 25) × 6 = 100 × 6 = 600，这就是乘法分配律的应用。'),
        ('6', 'B', '102 × 35 = (100 + 2) × 35 = 100 × 35 + 2 × 35，所以B正确。'),
    ]
    for num, ans, explain in ans_choice:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}. ')
        run.bold = True
        run = p.add_run(f'答案：{ans}')
        p = doc.add_paragraph()
        run = p.add_run(f'   {explain}')
        run.font.size = Pt(11)
        doc.add_paragraph()
    
    doc.add_paragraph()
    
    # 简便计算答案
    p = doc.add_heading('三、简便计算', level=1)
    ans_calc = [
        ('1.  98 × 45', 
         '解：98 × 45 = (100 - 2) × 45 = 100 × 45 - 2 × 45 = 4500 - 90 = 4410',
         '把98看成100-2，用乘法分配律展开，两项都要乘45，不要漏乘。'),
        ('2.  102 × 36',
         '解：102 × 36 = (100 + 2) × 36 = 100 × 36 + 2 × 36 = 3600 + 72 = 3672',
         '把102拆成100+2，乘法分配律，每一项都要乘36。'),
        ('3.  125 × 48',
         '解：125 × 48 = 125 × (8 × 6) = (125 × 8) × 6 = 1000 × 6 = 6000',
         '看到125想到找8，因为125×8=1000，这样计算简便。这里用乘法结合律，不需要分配。'),
        ('4.  25 × 125 × 32',
         '解：25 × 125 × 32 = 25 × 125 × (4 × 8) = (25 × 4) × (125 × 8) = 100 × 1000 = 100000',
         '看到25找4，看到125找8，所以把32拆成4×8，用乘法结合律凑整。'),
        ('5.  (30 + 4) × 25',
         '解：(30 + 4) × 25 = 30 × 25 + 4 × 25 = 750 + 100 = 850',
         '乘法分配律，括号里30和4两项都要和25相乘，不要漏乘。'),
        ('6.  456 - 56 - 44',
         '解：456 - 56 - 44 = 456 - (56 + 44) = 456 - 100 = 356',
         '一个数连续减两个数，等于这个数减这两个数的和。'
         '添括号时，括号前面是减号，括号里面要变号，所以 456 - 56 - 44 = 456 - (56 + 44)。'),
    ]
    for q, ans, explain in ans_calc:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.size = Pt(12)
        p = doc.add_paragraph()
        run = p.add_run(f'   {ans}')
        p = doc.add_paragraph()
        run = p.add_run(f'   {explain}')
        run.font.size = Pt(11)
        doc.add_paragraph()
    
    doc.add_paragraph()
    
    # 解决问题答案
    p = doc.add_heading('四、解决问题', level=1)
    p = doc.add_paragraph()
    run = p.add_run('解答：')
    p = doc.add_paragraph()
    run = p.add_run('1. 每条利润 = 售价 - 进价 = 135 - 85 = 50 元')
    p = doc.add_paragraph()
    run = p.add_run('2. 总条数 = 总利润 ÷ 每条利润 = 6000 ÷ 50 = 120 条')
    p = doc.add_paragraph()
    run = p.add_run('**答案：这批牛仔裤一共有 120 条**')
    p = doc.add_paragraph()
    run = p.add_run('   提示：本题没有多余干扰信息，直接计算即可。')
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 评分标准
    p = doc.add_heading('评分标准', level=1)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '总分'
    hdr[1].text = '评价'
    hdr[2].text = '建议'
    row = table.rows[1].cells
    row[0].text = '90-100\n70-89\n<70'
    row[1].text = '🎉 优秀\n👍 良好\n📚 需要加强'
    row[2].text = '整理错题，考前复习\n针对错题知识点再练\n回归课本，重新理解概念'
    
    doc.add_paragraph()
    
    # 末尾
    p = doc.add_paragraph()
    run = p.add_run('📖 复习提示：做错的题目对应知识点，请回到错题本对照复习')
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存
    output_path = '/home/wangfei/.openclaw/workspace/四年级数学错题专项复习卷第二套.docx'
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    output = create_fresh_test()
    print(f"Word文档已生成: {output}")