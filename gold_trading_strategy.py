#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_gold_trading_strategy():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 创建标题样式
    title_style = doc.styles.add_style('Title Style', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = '微软雅黑'
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)
    
    # 创建一级标题样式
    h1_style = doc.styles.add_style('Heading 1 Style', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = '微软雅黑'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(12)
    
    # 创建二级标题样式
    h2_style = doc.styles.add_style('Heading 2 Style', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = '微软雅黑'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(51, 102, 153)
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(8)
    
    # 创建三级标题样式
    h3_style = doc.styles.add_style('Heading 3 Style', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = '微软雅黑'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(70, 130, 180)
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(6)
    
    # 创建重点样式
    highlight_style = doc.styles.add_style('Highlight Style', WD_STYLE_TYPE.PARAGRAPH)
    highlight_style.font.name = '微软雅黑'
    highlight_style.font.size = Pt(11)
    highlight_style.font.bold = True
    highlight_style.font.color.rgb = RGBColor(204, 0, 0)
    highlight_style.paragraph_format.space_after = Pt(6)
    
    # 创建列表样式
    list_style = doc.styles.add_style('List Style', WD_STYLE_TYPE.PARAGRAPH)
    list_style.font.name = '微软雅黑'
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.3)
    list_style.paragraph_format.space_after = Pt(4)
    
    # 文档标题
    title = doc.add_paragraph('2026年黄金交易策略实战指南', style='Title Style')
    
    # 副标题
    subtitle = doc.add_paragraph('基本面分析 + 技术分析 + 风险管理 + 实战策略')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    subtitle.paragraph_format.space_after = Pt(30)
    
    # 目录
    doc.add_paragraph('目录', style='Heading 1 Style')
    toc_items = [
        '一、2026年黄金市场宏观分析',
        '二、影响黄金价格的核心因素',
        '三、技术分析体系详解',
        '四、日内交易策略',
        '五、波段交易策略',
        '六、长线投资策略',
        '七、风险管理体系',
        '八、实战交易模板',
        '九、常见交易误区',
        '十、交易心理建设'
    ]
    for item in toc_items:
        doc.add_paragraph(f'    {item}', style='List Style')
    
    doc.add_page_break()
    
    # 第一章
    doc.add_paragraph('一、2026年黄金市场宏观分析', style='Heading 1 Style')
    
    doc.add_paragraph('1.1 市场整体格局', style='Heading 2 Style')
    doc.add_paragraph('2026年黄金市场呈现"高位震荡、中枢上移"的特征。支撑金价的核心逻辑在于全球央行持续购金（去美元化趋势）与美联储降息周期（实际利率下行）的共振效应。尽管短期可能因获利了结出现技术性回调，但长期上涨趋势未改。')
    
    doc.add_paragraph('1.2 机构预测汇总', style='Heading 2 Style')
    doc.add_paragraph('• 高盛：2026年底目标价 $5,200-5,500/盎司', style='List Style')
    doc.add_paragraph('• 摩根大通：维持超配评级，目标价 $5,300/盎司', style='List Style')
    doc.add_paragraph('• 花旗银行：预计全年均价 $4,800/盎司', style='List Style')
    doc.add_paragraph('• 瑞银：看涨至 $5,600/盎司，推荐战略性配置', style='List Style')
    
    doc.add_paragraph('1.3 季度走势预测', style='Heading 2 Style')
    doc.add_paragraph('• Q1（1-3月）：高位震荡，消化前期涨幅，支撑位 $4,500', style='List Style')
    doc.add_paragraph('• Q2（4-6月）：震荡上行，美联储降息预期升温，目标 $4,900', style='List Style')
    doc.add_paragraph('• Q3（7-9月）：主升浪开启，央行购金旺季，目标 $5,200', style='List Style')
    doc.add_paragraph('• Q4（10-12月）：冲高回落，年底获利了结，支撑位 $4,800', style='List Style')
    
    doc.add_paragraph('⚠️ 风险提示：地缘政治冲突升级、美联储政策转向、美元指数大幅波动是主要风险点。', style='Highlight Style')
    
    # 第二章
    doc.add_paragraph('二、影响黄金价格的核心因素', style='Heading 1 Style')
    
    doc.add_paragraph('2.1 货币政策与实际利率', style='Heading 2 Style')
    doc.add_paragraph('黄金价格与实际利率呈强负相关。当美联储进入降息周期，实际利率下行，黄金的持有成本降低，价格通常上涨。')
    doc.add_paragraph('• 关注美联储FOMC会议（每年8次）', style='List Style')
    doc.add_paragraph('• CPI、PCE通胀数据', style='List Style')
    doc.add_paragraph('• 非农就业报告', style='List Style')
    doc.add_paragraph('• 10年期美债收益率', style='List Style')
    
    doc.add_paragraph('2.2 美元指数', style='Heading 2 Style')
    doc.add_paragraph('黄金以美元计价，与美元指数通常呈负相关关系。美元走弱时，以其他货币计价的黄金相对便宜，需求增加。')
    
    doc.add_paragraph('2.3 央行购金趋势', style='Heading 2 Style')
    doc.add_paragraph('全球央行持续增持黄金储备，特别是新兴市场国家的去美元化进程，为金价提供长期支撑。2025年全球央行购金量创历史新高，2026年预计维持强劲势头。')
    
    doc.add_paragraph('2.4 地缘政治风险', style='Heading 2 Style')
    doc.add_paragraph('黄金作为传统避险资产，在地缘冲突、战争风险升温时往往获得避险买盘。')
    
    doc.add_paragraph('2.5 通胀预期', style='Heading 2 Style')
    doc.add_paragraph('黄金被视为对抗通胀的硬通货。当通胀预期升温时，黄金的保值需求增加。')
    
    # 第三章
    doc.add_paragraph('三、技术分析体系详解', style='Heading 1 Style')
    
    doc.add_paragraph('3.1 多周期分析框架', style='Heading 2 Style')
    doc.add_paragraph('• 日线图（D1）：判断趋势方向和关键支撑阻力位', style='List Style')
    doc.add_paragraph('• 4小时图（H4）：寻找交易机会和入场时机', style='List Style')
    doc.add_paragraph('• 1小时图（H1）：精确入场点位和止损设置', style='List Style')
    doc.add_paragraph('• 15分钟图（M15）：超短线交易的精细化操作', style='List Style')
    
    doc.add_paragraph('3.2 核心技术指标', style='Heading 2 Style')
    
    doc.add_paragraph('3.2.1 移动平均线（MA）', style='Heading 3 Style')
    doc.add_paragraph('• MA5/MA10：短期趋势判断，金叉做多，死叉做空', style='List Style')
    doc.add_paragraph('• MA20/MA60：中期趋势，MA60作为多空分水岭', style='List Style')
    doc.add_paragraph('• MA200：长期趋势线，价格在上方为牛市，下方为熊市', style='List Style')
    
    doc.add_paragraph('3.2.2 相对强弱指标（RSI）', style='Heading 3 Style')
    doc.add_paragraph('• RSI > 70：超买区域，警惕回调风险', style='List Style')
    doc.add_paragraph('• RSI < 30：超卖区域，关注反弹机会', style='List Style')
    doc.add_paragraph('• RSI 背离：趋势可能反转的重要信号', style='List Style')
    
    doc.add_paragraph('3.2.3 MACD指标', style='Heading 3 Style')
    doc.add_paragraph('• 金叉：DIF上穿DEA，多头信号', style='List Style')
    doc.add_paragraph('• 死叉：DIF下穿DEA，空头信号', style='List Style')
    doc.add_paragraph('• 顶背离/底背离：趋势反转预警', style='List Style')
    
    doc.add_paragraph('3.2.4 布林带（Bollinger Bands）', style='Heading 3 Style')
    doc.add_paragraph('• 上轨：压力位，触及可考虑做空', style='List Style')
    doc.add_paragraph('• 中轨：MA20，趋势方向判断', style='List Style')
    doc.add_paragraph('• 下轨：支撑位，触及可考虑做多', style='List Style')
    
    doc.add_paragraph('3.3 K线形态识别', style='Heading 2 Style')
    
    doc.add_paragraph('反转形态：', style='Heading 3 Style')
    doc.add_paragraph('• 锤子线/上吊线：底部/顶部反转信号', style='List Style')
    doc.add_paragraph('• 吞没形态：强烈的反转信号', style='List Style')
    doc.add_paragraph('• 晨星/暮星：三天反转形态，可靠性高', style='List Style')
    doc.add_paragraph('• 头肩顶/头肩底：经典趋势反转形态', style='List Style')
    doc.add_paragraph('• 双顶/双底：M头/W底形态', style='List Style')
    
    doc.add_paragraph('持续形态：', style='Heading 3 Style')
    doc.add_paragraph('• 三角形整理：上升三角形看涨，下降三角形看跌', style='List Style')
    doc.add_paragraph('• 旗形/三角旗形：趋势持续信号', style='List Style')
    doc.add_paragraph('• 矩形整理：箱体震荡，突破后沿趋势方向运行', style='List Style')
    
    doc.add_paragraph('3.4 支撑与阻力位判断', style='Heading 2 Style')
    doc.add_paragraph('• 历史高低点：前期重要的价格高点和低点', style='List Style')
    doc.add_paragraph('• 整数关口：如 $4,500、$5,000 等心理关口', style='List Style')
    doc.add_paragraph('• 斐波那契回撤位：23.6%、38.2%、50%、61.8%', style='List Style')
    doc.add_paragraph('• 趋势线：连接高点或低点形成的趋势通道', style='List Style')
    
    # 第四章
    doc.add_paragraph('四、日内交易策略', style='Heading 1 Style')
    
    doc.add_paragraph('4.1 日内交易的核心原则', style='Heading 2 Style')
    doc.add_paragraph('日内交易追求"高频机会+不留隔夜风险"，黄金市场波动充足、流动性极强，是日内交易的理想标的。但要求严格的执行力和纪律性。')
    
    doc.add_paragraph('4.2 最佳交易时段', style='Heading 2 Style')
    doc.add_paragraph('• 亚盘（7:00-14:00）：波动相对较小，适合观察和准备', style='List Style')
    doc.add_paragraph('• 欧盘（14:00-20:30）：波动加大，开始出现交易机会', style='List Style')
    doc.add_paragraph('• 美盘（20:30-5:00）：波动最大，数据公布密集，主交易时段', style='List Style')
    doc.add_paragraph('⭐ 重点关注：20:30-22:00 美国数据公布时段，波动最为剧烈', style='Highlight Style')
    
    doc.add_paragraph('4.3 突破交易策略', style='Heading 2 Style')
    doc.add_paragraph('策略逻辑：价格突破关键支撑阻力位后跟进交易')
    doc.add_paragraph('• 入场条件：实体K线突破关键价位并站稳', style='List Style')
    doc.add_paragraph('• 止损设置：突破位反向10-15美元', style='List Style')
    doc.add_paragraph('• 止盈目标：1:2或1:3的盈亏比', style='List Style')
    doc.add_paragraph('• 过滤条件：配合成交量放大确认', style='List Style')
    
    doc.add_paragraph('4.4 回调交易策略', style='Heading 2 Style')
    doc.add_paragraph('策略逻辑：趋势确立后，回调至支撑阻力位入场')
    doc.add_paragraph('• 入场条件：回调至MA20/MA60或斐波那契回撤位', style='List Style')
    doc.add_paragraph('• 止损设置：支撑阻力位外5-10美元', style='List Style')
    doc.add_paragraph('• 止盈目标：前期高低点或整数关口', style='List Style')
    doc.add_paragraph('• 优势：盈亏比高，胜率相对稳定', style='List Style')
    
    doc.add_paragraph('4.5 区间交易策略', style='Heading 2 Style')
    doc.add_paragraph('策略逻辑：在震荡区间内高抛低吸')
    doc.add_paragraph('• 入场条件：触及区间上轨做空，触及下轨做多', style='List Style')
    doc.add_paragraph('• 止损设置：区间外5-8美元', style='List Style')
    doc.add_paragraph('• 止盈目标：区间中部或对侧轨道', style='List Style')
    doc.add_paragraph('• 注意：一旦区间突破，立即止损并转向突破策略', style='List Style')
    
    doc.add_paragraph('4.6 新闻交易策略', style='Heading 2 Style')
    doc.add_paragraph('策略逻辑：重要数据公布后跟随市场情绪交易')
    doc.add_paragraph('• 关注数据：非农、CPI、美联储决议、GDP', style='List Style')
    doc.add_paragraph('• 操作方法：数据公布后等待15分钟，趋势确立后跟进', style='List Style')
    doc.add_paragraph('• 风险提示：数据行情波动极大，建议轻仓操作', style='List Style')
    
    # 第五章
    doc.add_paragraph('五、波段交易策略', style='Heading 1 Style')
    
    doc.add_paragraph('5.1 波段交易的特点', style='Heading 2 Style')
    doc.add_paragraph('波段交易持仓时间通常为3-15个交易日，追求捕捉中期趋势行情，交易频率较低但单笔盈利空间较大，适合时间相对有限的交易者。')
    
    doc.add_paragraph('5.2 趋势跟踪策略', style='Heading 2 Style')
    doc.add_paragraph('• 趋势确认：MA60方向判断大趋势，价格在MA60上方做多，下方做空', style='List Style')
    doc.add_paragraph('• 入场信号：MA5金叉MA20，配合RSI在50以上', style='List Style')
    doc.add_paragraph('• 止损设置：最近的波段低点下方15-20美元', style='List Style')
    doc.add_paragraph('• 止盈方式：移动止损追踪趋势，或达到前期重要阻力位', style='List Style')
    
    doc.add_paragraph('5.3 形态交易策略', style='Heading 2 Style')
    doc.add_paragraph('重点关注日线级别的经典形态：')
    doc.add_paragraph('• 头肩底/头肩顶：形态完成后入场，目标为形态高度', style='List Style')
    doc.add_paragraph('• 双底/双顶：颈线突破后入场，目标为形态高度', style='List Style')
    doc.add_paragraph('• 三角形突破：三角形末端突破后跟进', style='List Style')
    doc.add_paragraph('• 箱体突破：突破箱体上下沿后跟进交易', style='List Style')
    
    doc.add_paragraph('5.4 周线级别策略', style='Heading 2 Style')
    doc.add_paragraph('适合大资金配置，持仓周期1-3个月：')
    doc.add_paragraph('• 周线MA200判断长期趋势', style='List Style')
    doc.add_paragraph('• 周线RSI低于30关注做多机会，高于70警惕见顶', style='List Style')
    doc.add_paragraph('• 周线级别K线形态反转信号可靠性更高', style='List Style')
    doc.add_paragraph('• 建议仓位：总资金的10%-20%', style='List Style')
    
    # 第六章
    doc.add_paragraph('六、长线投资策略', style='Heading 1 Style')
    
    doc.add_paragraph('6.1 战略配置思路', style='Heading 2 Style')
    doc.add_paragraph('黄金作为资产配置的"压舱石"，建议配置比例为总资产的5%-15%。2026年在美联储降息周期和央行购金的双重驱动下，可适当提高配置比例至10%-20%。')
    
    doc.add_paragraph('6.2 定投策略', style='Heading 2 Style')
    doc.add_paragraph('• 定投方式：每月固定金额买入，忽略短期波动', style='List Style')
    doc.add_paragraph('• 适合人群：普通投资者，无需盯盘', style='List Style')
    doc.add_paragraph('• 投资标的：实物黄金、黄金ETF、纸黄金', style='List Style')
    doc.add_paragraph('• 优势：摊薄成本，长期复利效应', style='List Style')
    
    doc.add_paragraph('6.3 逆向投资策略', style='Heading 2 Style')
    doc.add_paragraph('• 在市场恐慌、金价大幅回调时分批建仓', style='List Style')
    doc.add_paragraph('• 在市场疯狂、多头情绪极端时分批止盈', style='List Style')
    doc.add_paragraph('• 需要强大的耐心和逆向思维能力', style='List Style')
    
    # 第七章
    doc.add_paragraph('七、风险管理体系', style='Heading 1 Style')
    
    doc.add_paragraph('7.1 仓位管理原则', style='Heading 2 Style')
    doc.add_paragraph('⭐ 黄金法则：任何单笔交易的亏损不超过总资金的1%-2%', style='Highlight Style')
    doc.add_paragraph('• 日内交易：单笔仓位不超过总资金的5%', style='List Style')
    doc.add_paragraph('• 波段交易：单笔仓位不超过总资金的10%', style='List Style')
    doc.add_paragraph('• 长线投资：单笔仓位不超过总资金的20%', style='List Style')
    doc.add_paragraph('• 总敞口：任何时候总仓位不超过总资金的50%', style='List Style')
    
    doc.add_paragraph('7.2 止损设置方法', style='Heading 2 Style')
    doc.add_paragraph('• 固定止损：入场前设定固定止损金额', style='List Style')
    doc.add_paragraph('• 技术止损：设置在关键支撑阻力位外侧', style='List Style')
    doc.add_paragraph('• 移动止损：盈利后逐步上移止损，锁定利润', style='List Style')
    doc.add_paragraph('• 时间止损：持仓一定时间未达预期则平仓', style='List Style')
    doc.add_paragraph('⚠️ 重要：止损必须在入场前设定，严格执行！', style='Highlight Style')
    
    doc.add_paragraph('7.3 止盈策略', style='Heading 2 Style')
    doc.add_paragraph('• 目标止盈：设定固定盈利目标', style='List Style')
    doc.add_paragraph('• 分批止盈：达到第一目标平半仓，剩余持仓移动止损', style='List Style')
    doc.add_paragraph('• 技术止盈：遇到关键支撑阻力位止盈', style='List Style')
    doc.add_paragraph('• 追踪止盈：使用移动止损锁定趋势利润', style='List Style')
    
    doc.add_paragraph('7.4 风险控制 Checklist', style='Heading 2 Style')
    doc.add_paragraph('交易前检查：')
    doc.add_paragraph('• 单笔风险是否控制在2%以内？', style='List Style')
    doc.add_paragraph('• 止损位是否合理并已设置？', style='List Style')
    doc.add_paragraph('• 盈亏比是否大于1:1.5？', style='List Style')
    doc.add_paragraph('• 总仓位是否在50%以下？', style='List Style')
    doc.add_paragraph('• 是否避开重大数据公布前15分钟？', style='List Style')
    
    # 第八章
    doc.add_paragraph('八、实战交易模板', style='Heading 1 Style')
    
    doc.add_paragraph('8.1 日内交易实战模板', style='Heading 2 Style')
    
    doc.add_paragraph('模板一：突破交易', style='Heading 3 Style')
    doc.add_paragraph('时间框架：H1 + M15')
    doc.add_paragraph('1. H1图识别关键支撑阻力位', style='List Style')
    doc.add_paragraph('2. M15图等待价格测试该位置', style='List Style')
    doc.add_paragraph('3. 实体K线突破后入场（收盘价确认）', style='List Style')
    doc.add_paragraph('4. 止损：突破位反向10美元', style='List Style')
    doc.add_paragraph('5. 止盈：第一目标15美元，第二目标30美元', style='List Style')
    doc.add_paragraph('6. 盈亏比：1:1.5 ~ 1:3', style='List Style')
    
    doc.add_paragraph('模板二：回调做多（上升趋势）', style='Heading 3 Style')
    doc.add_paragraph('时间框架：H4 + H1')
    doc.add_paragraph('1. H4图确认上升趋势（MA20向上，价格在MA20上方）', style='List Style')
    doc.add_paragraph('2. H1图等待价格回调至MA20或38.2%回撤位', style='List Style')
    doc.add_paragraph('3. 出现看涨K线形态（锤子线、吞没等）后入场', style='List Style')
    doc.add_paragraph('4. 止损：回调低点下方8-10美元', style='List Style')
    doc.add_paragraph('5. 止盈：前期高点或下一个阻力位', style='List Style')
    
    doc.add_paragraph('8.2 波段交易实战模板', style='Heading 2 Style')
    
    doc.add_paragraph('模板：趋势跟踪', style='Heading 3 Style')
    doc.add_paragraph('时间框架：D1 + H4')
    doc.add_paragraph('1. D1图确认趋势方向（MA60作为判断标准）', style='List Style')
    doc.add_paragraph('2. D1图MA5金叉MA20作为入场信号', style='List Style')
    doc.add_paragraph('3. H4图寻找精确入场点（回调至支撑位）', style='List Style')
    doc.add_paragraph('4. 止损：最近波段低点下方15-20美元', style='List Style')
    doc.add_paragraph('5. 止盈：使用移动止损（MA20跟踪）', style='List Style')
    doc.add_paragraph('6. 持仓周期：3-15个交易日', style='List Style')
    
    # 第九章
    doc.add_paragraph('九、常见交易误区', style='Heading 1 Style')
    
    doc.add_paragraph('9.1 技术分析误区', style='Heading 2 Style')
    doc.add_paragraph('• ❌ 过度使用指标：指标不是越多越好，2-3个核心指标足够', style='List Style')
    doc.add_paragraph('• ❌ 忽视大周期：小周期服从大周期，顺大趋势交易', style='List Style')
    doc.add_paragraph('• ❌ 预测顶底：不要猜顶猜底，跟随趋势', style='List Style')
    doc.add_paragraph('• ❌ 忽略成交量：成交量验证价格突破的有效性', style='List Style')
    
    doc.add_paragraph('9.2 风险管理误区', style='Heading 2 Style')
    doc.add_paragraph('• ❌ 不止损：扛单是爆仓的主要原因', style='List Style')
    doc.add_paragraph('• ❌ 频繁止损：止损位设置不合理，被来回扫', style='List Style')
    doc.add_paragraph('• ❌ 重仓交易：追求短期暴利，忽视风险', style='List Style')
    doc.add_paragraph('• ❌ 亏损加仓：摊平成本可能导致更大亏损', style='List Style')
    
    doc.add_paragraph('9.3 交易心理误区', style='Heading 2 Style')
    doc.add_paragraph('• ❌ 贪婪：盈利后不愿止盈，最终回吐利润', style='List Style')
    doc.add_paragraph('• ❌ 恐惧：亏损时不敢止损，越扛越大', style='List Style')
    doc.add_paragraph('• ❌ 报复性交易：连续亏损后急于翻本，越做越错', style='List Style')
    doc.add_paragraph('• ❌ 过度自信：连续盈利后盲目自大，放松风控', style='List Style')
    
    # 第十章
    doc.add_paragraph('十、交易心理建设', style='Heading 1 Style')
    
    doc.add_paragraph('10.1 建立交易系统', style='Heading 2 Style')
    doc.add_paragraph('• 明确入场条件、止损设置、止盈策略', style='List Style')
    doc.add_paragraph('• 制定仓位管理规则', style='List Style')
    doc.add_paragraph('• 写下交易计划并严格执行', style='List Style')
    doc.add_paragraph('• 系统简单可执行，避免复杂化', style='List Style')
    
    doc.add_paragraph('10.2 交易日志', style='Heading 2 Style')
    doc.add_paragraph('每笔交易记录：')
    doc.add_paragraph('• 入场理由、止损止盈设置', style='List Style')
    doc.add_paragraph('• 实际出场点和盈亏情况', style='List Style')
    doc.add_paragraph('• 交易过程中的情绪变化', style='List Style')
    doc.add_paragraph('• 总结经验教训', style='List Style')
    doc.add_paragraph('定期（每周/每月）复盘，持续优化交易系统。', style='List Style')
    
    doc.add_paragraph('10.3 心态修炼', style='Heading 2 Style')
    doc.add_paragraph('• 接受亏损是交易的一部分，胜率50%-60%已很优秀', style='List Style')
    doc.add_paragraph('• 保持耐心，等待符合系统的交易机会', style='List Style')
    doc.add_paragraph('• 连续亏损时暂停交易，冷静分析原因', style='List Style')
    doc.add_paragraph('• 不要与他人比较，专注自己的成长', style='List Style')
    doc.add_paragraph('• 交易之外有生活，保持身心平衡', style='List Style')
    
    # 结语
    doc.add_paragraph('结语', style='Heading 1 Style')
    doc.add_paragraph('黄金交易是一场马拉松，不是百米冲刺。成功的交易者不是靠一次两次的暴利，而是靠持续稳定的盈利和严格的风险管理。')
    doc.add_paragraph('记住：')
    doc.add_paragraph('• 趋势为王，顺势而为', style='List Style')
    doc.add_paragraph('• 止损是生命线，严格执行', style='List Style')
    doc.add_paragraph('• 仓位管理决定你能走多远', style='List Style')
    doc.add_paragraph('• 持续学习，不断进化', style='List Style')
    doc.add_paragraph('')
    doc.add_paragraph('祝您在2026年黄金市场中取得理想的投资回报！💰', style='Highlight Style')
    
    # 保存文档
    doc.save('/home/wangfei/.openclaw/workspace/2026年黄金交易策略实战指南.docx')
    print("文档创建成功：2026年黄金交易策略实战指南.docx")

if __name__ == '__main__':
    create_gold_trading_strategy()
