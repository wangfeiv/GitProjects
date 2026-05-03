#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

OUT = '/home/wangfei/.openclaw/workspace/多多四年级数学错题专项复习卷-重制版.docx'


def set_font(run, size=12, bold=False):
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(size)
    run.bold = bold


def add_p(doc, text='', size=12, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def add_lines(doc, lines, size=12):
    for line in lines:
        add_p(doc, line, size=size)


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)

    add_p(doc, '多多四年级数学错题专项复习卷（重制版）', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, '姓名：__________    日期：__________    用时：__________    得分：__________', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, '说明：本卷根据错题本中的薄弱点重新命题，全部为纯文本，不含特殊公式字符。')
    add_p(doc, '覆盖知识点：小数的组成、四舍五入求原数、用字母表示数、路程公式、长方体最大最小面、乘法分配律、乘法结合律、添括号变号、利润问题。')

    add_p(doc, '一、填空题（每空 4 分，共 32 分）', bold=True)
    add_lines(doc, [
        '1. 用 1、0、4、7 四个数字和小数点组成数（每个数字都要用一次）。',
        '   最大的一位小数是____________，最小的两位小数是____________。',
        '2. 一个三位小数四舍五入后是 8.60，这个三位小数最大是____________，最小是____________。',
        '3. 一辆汽车每小时行 v 千米，行了 t 小时，一共行了____________千米。',
        '   如果路程用 s 表示，公式是____________。',
        '4. 一个长方体，长 12 分米，宽 9 分米，高 4 分米。',
        '   最大一个面的面积是____________平方分米，最小一个面的面积是____________平方分米。',
        '5. 小明每分钟走 a 米，8 分钟走____________米；如果一共走了 240 米，需要____________分钟。',
    ])

    add_p(doc, '二、选择题（每题 4 分，共 24 分）', bold=True)
    add_lines(doc, [
        '6. 甲数是 a，比乙数的 6 倍多 c，乙数是（    ）。',
        '   A. 6a+c    B. 6a-c    C. (a+c)÷6    D. (a-c)÷6',
        '7. 小强跑 100 米用了 x 秒，小明比小强快 3 秒，小明用了（    ）秒。',
        '   A. x+3    B. x-3    C. 3x    D. x÷3',
        '8. 苹果每千克 x 元，梨每千克 y 元，买 4 千克苹果比买 3 千克梨便宜多少元？正确列式是（    ）。',
        '   A. 4x-3y    B. 3y-4x    C. 4x+3y    D. (4-3)xy',
        '9. 聪聪计算 (4+□)×30 时，错算成了 4+□×30，错误结果与正确结果相差（    ）。',
        '   A. 26    B. 116    C. 120    D. 4',
        '10. 下面适合用乘法结合律简便计算的是（    ）。',
        '   A. (125+8)×4    B. 25×(40×4)    C. (80+2)×35    D. 36×99',
        '11. 计算 280-80+25 时，正确的是（    ）。',
        '   A. 280-(80+25)    B. 280-(80-25)    C. (280-80)+25    D. A 和 C 都对',
    ])

    add_p(doc, '三、用简便方法计算（每题 6 分，共 36 分）', bold=True)
    add_lines(doc, [
        '12. 103×26',
        '13. 98×45',
        '14. 25×(40×4)',
        '15. 125×32×25',
        '16. (125+16)×4',
        '17. 356-56-44',
    ])

    add_p(doc, '四、解决问题（每题 4 分，共 8 分）', bold=True)
    add_lines(doc, [
        '18. 商店运进一批外套，进价每件 48 元，售价每件 65 元。全部卖完一共赚了 408 元。',
        '    这批外套一共有多少件？',
        '19. 水果店运来苹果和香蕉各 9 箱。苹果每箱 35 千克，香蕉每箱 45 千克。',
        '    一共运来多少千克水果？请用简便方法解答。',
    ])

    doc.add_page_break()

    add_p(doc, '参考答案与解析', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_p(doc, '一、填空题答案', bold=True)
    add_lines(doc, [
        '1. 最大的一位小数：741.0    最小的两位小数：10.47',
        '   解析：一位小数表示小数点后只有 1 位；两位小数表示小数点后有 2 位。',
        '2. 最大：8.604    最小：8.595',
        '   解析：最大是“四舍”得到，所以最后一位最大是 4；最小是“五入”得到，所以最后一位最小是 5。',
        '3. vt；s=vt',
        '4. 最大面：12×9=108    最小面：9×4=36',
        '5. 8a；240÷a',
    ])

    add_p(doc, '二、选择题答案', bold=True)
    add_lines(doc, [
        '6. D',
        '   解析：a=6×乙+c，所以 6×乙=a-c，乙=(a-c)÷6。',
        '7. B',
        '   解析：“快”表示用时更少，所以用减法。',
        '8. B',
        '   解析：“苹果比梨便宜多少”就是 梨的钱数 - 苹果的钱数。',
        '9. B',
        '   解析：正确结果比错误结果多了 4×30-4=120-4=116。',
        '10. B',
        '   解析：括号里是乘法，适合用乘法结合律。',
        '11. D',
        '   解析：280-80+25 可以按顺序算，也可以写成 280-(80-25)。不能写成 280-(80+25)。',
    ])

    add_p(doc, '三、简便计算答案', bold=True)
    add_lines(doc, [
        '12. 103×26=(100+3)×26=100×26+3×26=2600+78=2678',
        '13. 98×45=(100-2)×45=100×45-2×45=4500-90=4410',
        '14. 25×(40×4)=(25×4)×40=100×40=4000',
        '15. 125×32×25=125×(8×4)×25=(125×8)×(4×25)=1000×100=100000',
        '16. (125+16)×4=125×4+16×4=500+64=564',
        '17. 356-56-44=356-(56+44)=356-100=256',
    ])

    add_p(doc, '四、解决问题答案', bold=True)
    add_lines(doc, [
        '18. 每件利润：65-48=17（元）',
        '    一共有：408÷17=24（件）',
        '    答：这批外套一共有 24 件。',
        '19. 方法一：35×9+45×9=(35+45)×9=80×9=720（千克）',
        '    答：一共运来 720 千克水果。',
    ])

    add_p(doc, '出卷时间：' + datetime.now().strftime('%Y-%m-%d %H:%M'))
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build_doc()
