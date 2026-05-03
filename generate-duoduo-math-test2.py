#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成多多四年级数学错题专项复习卷(第二套)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def create_math_test():
    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(1.8)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 标题
    title = doc.add_heading('四年级数学下册\n错题专项复习卷(第二套)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True

    # 信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('姓名:___________    用时:___________    得分:___________')
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph('📋 本卷覆盖知识点:小数组合、四舍五入、用字母表示数、乘法分配律、乘法结合律、添括号变号、长方体面积。')
    doc.add_paragraph()

    # 第一部分:填空题
    doc.add_heading('一、填空题(每空3分,共36分)', level=1)

    p = doc.add_paragraph()
    run = p.add_run('1. 用2、0、5、7四个数字和小数点组成(每个数字都要用,只能用一次):')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('   - 最大的一位小数是 ____________')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小的两位小数是 ____________')

    p = doc.add_paragraph()
    run = p.add_run('2. 一个两位小数四舍五入后是8.6,这个两位小数:')
    p = doc.add_paragraph()
    run = p.add_run('   - 最大是 ____________')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小是 ____________')

    p = doc.add_paragraph()
    run = p.add_run('3. 一个三位小数四舍五入后是5.00,这个三位小数:')
    p = doc.add_paragraph()
    run = p.add_run('   - 最大是 ____________')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小是 ____________')

    p = doc.add_paragraph()
    run = p.add_run('4. 一块长方体木块,长12厘米,宽8厘米,高5厘米。')
    p = doc.add_paragraph()
    run = p.add_run('   - 最大一个面的面积是 ____________ 平方厘米')
    p = doc.add_paragraph()
    run = p.add_run('   - 最小一个面的面积是 ____________ 平方厘米')

    p = doc.add_paragraph()
    run = p.add_run('5. 一辆货车每小时行驶x千米,行驶了y千米,需要 ____________ 小时。(用字母表示)')

    p = doc.add_paragraph()
    run = p.add_run('6. 苹果每千克a元,梨每千克b元。买5千克苹果比买3千克梨便宜 ____________ 元。(用字母表示)')

    p = doc.add_paragraph()
    run = p.add_run('7. 如果每分钟走m米,5分钟走 ____________ 米,走100米需要 ____________ 分钟。')

    doc.add_paragraph()

    # 第二部分:选择题
    doc.add_heading('二、选择题(每题4分,共24分)', level=1)

    questions_choice = [
        ('1. 甲数是x,比乙数的4倍少y,乙数是(  )',
         ['A. 4x - y', 'B. (x + y) ÷ 4', 'C. (x - y) ÷ 4', 'D. x ÷ 4 - y'],
         ''),
        ('2. 百米赛跑,小强用了a秒,小刚比小强慢3秒,小刚用了(  )秒',
         ['A. a + 3', 'B. a - 3', 'C. 3a', 'D. a ÷ 3'],
         ''),
        ('3. 长方形菜地长a米,宽比长少b米,面积是(  )平方米',
         ['A. a × b', 'B. a × (a + b)', 'C. a × (a - b)', 'D. 2a + 2(a - b)'],
         ''),
        ('4. 水果店运来苹果和梨各10筐,苹果每筐25千克,梨每筐35千克。一共运来多少千克?用简便方法计算,列式正确的是(  )',
         ['A. 25 × 10 + 35 × 10', 'B. 10 × 25 × 35', 'C. (25 + 35) × 10', 'D. 25 + 35 × 10'],
         ''),
        ('5. 小明计算 25 × (4 + 8) 时,错算成 25 × 4 + 8,他算出的结果和正确结果相差(  )',
         ['A. 100', 'B. 192', 'C. 200', 'D. 208'],
         ''),
        ('6. 与 201 × 75 计算结果不相等的式子是(  )',
         ['A. (200 + 1) × 75', 'B. 200 × 75 + 1 × 75', 'C. 200 × 75 + 75', 'D. 200 × 75 + 1'],
         ''),
    ]

    for num, question, options in questions_choice:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}')
        run.font.size = Pt(12)
        for opt in options:
            p = doc.add_paragraph()
            run = p.add_run(f'   {opt}')
            run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run('你的选择:____')
        run.font.size = Pt(11)
        doc.add_paragraph()

    # 第三部分：简便计算
    doc.add_heading('三、用简便方法计算（每题5分，共30分）', level=1)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('💡 提示：')
    run.bold = True
    run = p.add_run('想一想，怎么拆数才能凑整？什么时候用乘法分配律，什么时候用乘法结合律？')
    run.font.size = Pt(11)
    doc.add_paragraph()
    calc_questions = [
        '1.  \\( 101 \\times 65 \\quad\\quad \\text{（提示：把101拆成}100+1\\text{，用乘法分配律}\\right)$',
        '2.  \\( 99 \\times 42 \\quad\\quad \\text{（提示：把99看成}100-1\\text{，用乘法分配律}\\right)$',
        '3.  \\( 25 \\times 36 \\quad\\quad \\text{（提示：看到25，想到拆出一个4来凑100}\\right)$',
        '4.  \\( 125 \\times 24 \\times 25 \\quad\\quad \\text{（提示：看到125和25，把24拆成}8\\times 3\\text{，这样}125\\times 8=1000\\text{，}25\\text{呢？}\\right)$',
        '5.  \\( 25 \\times (40 + 8) \\quad\\quad \\text{（提示：括号里两个数，都要和25相乘，不要漏乘}\\right)$',
        '6.  \\( 368 - 68 - 32 \\quad\\quad \\text{（提示：连续减两个数，等于减这两个数的}\\textit{和}\\right)$',
    ]
    for q in calc_questions:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.size = Pt(12)
        p = doc.add_paragraph('\n')

    doc.add_paragraph()

    # 第四部分:解决问题
    doc.add_heading('四、解决问题(10分)', level=1)
    p = doc.add_paragraph()
    run = p.add_run('商店运进一批衬衫,进价每件45元,售价每件68元。全部卖完后总利润是920元。这批衬衫一共有多少件?(先算每件赚多少元,再算总件数)')
    run.font.size = Pt(12)
    p = doc.add_paragraph('\n\n\n\n')

    # 分页放答案
    doc.add_page_break()

    # 答案部分
    ans_title = doc.add_heading('参考答案与详细解析', 0)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 填空题答案
    p = doc.add_heading('一、填空题', level=1)
    ans_fill = [
        ('1', '最大一位小数:\\(\\mathbf{752.0}\\),最小两位小数:\\(\\mathbf{20.57}\\)',
         '解析:一位小数,小数点后1位,剩下三位放整数部分,从大到小排列;两位小数,最小不能以0开头,所以2开头,从小到大排列。'),
        ('2', '最大:\\(\\mathbf{8.64}\\),最小:\\(\\mathbf{8.55}\\)',
         '解析:四舍得到8.6 → 最大第二位是4;五入得到8.6 → 最小第二位是5。记住:最大四舍,最小五入。'),
        ('3', '最大:\\(\\mathbf{5.004}\\),最小:\\(\\mathbf{4.995}\\)',
         '同理,四舍五入规则:最大四舍,最小五入。5.005五入后变成5.01,所以最大只能是5.004。'),
        ('4', '最大面:\\(\\mathbf{12 × 8 = 96}\\) → 96 平方厘米;最小面:\\(\\mathbf{5 × 8 = 40}\\) → 40 平方厘米',
         '规律:最大面 = 最长两条边相乘,最小面 = 最短两条边相乘。'),
        ('5', '答案:\\(\\mathbf{y \\div x}\\)',
         '公式:时间 = 路程 ÷ 速度 → y ÷ x。'),
        ('6', '答案:\\(\\mathbf{3b - 5a}\\)',
         '"A比B便宜多少" = B总价 - A总价 → 3b - 5a。注意谁比谁,用大减小。'),
        ('7', '答案:5分钟走 \\(\\mathbf{5m}\\) 米,走100米需要 \\(\\mathbf{100 \\div m}\\) 分钟',
         '路程 = 速度 × 时间;时间 = 路程 ÷ 速度。'),
    ]
    for num, ans, explain in ans_fill:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}. ')
        run.bold = True
        run = p.add_run(ans)
        p = doc.add_paragraph()
        run = p.add_run(f'   {explain}')
        run.font.size = Pt(11)
        doc.add_paragraph()

    doc.add_paragraph()

    # 选择题答案
    p = doc.add_heading('二、选择题', level=1)
    ans_choice = [
        ('1', 'B', '推导:甲数 = 乙数 × 4 - y → x = 4乙 - y → 4乙 = x + y → 乙 = (x + y) ÷ 4'),
        ('2', 'A', '"慢" → 用时更多 → a + 3。快用减法,慢用加法。'),
        ('3', 'C', '宽 = 长 - b = a - b → 面积 = 长 × 宽 = a × (a - b)'),
        ('4', 'C', '(25 + 35) × 10 = 60 × 10 = 600,这就是乘法分配律的应用。'),
        ('5', 'B', '正确结果:25×(4+8)=25×12=300,错误结果:25×4+8=108,相差:300-108=192'),
        ('6', 'D', '201×75=(200+1)×75=200×75+1×75=200×75+75,所以D少乘了,不相等。'),
    ]
    for num, ans, explain in ans_choice:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}. ')
        run.bold = True
        run = p.add_run(f'答案:{ans}')
        p = doc.add_paragraph()
        run = p.add_run(f'   {explain}')
        run.font.size = Pt(11)
        doc.add_paragraph()

    doc.add_paragraph()

    # 简便计算答案
    p = doc.add_heading('三、简便计算', level=1)
    ans_calc = [
        ('1.  \\( 101 \\times 65 \\)',
         '\\[\n\\begin{aligned}\n101 \\times 65 &= (100 + 1) \\times 65 \\\\\n&= 100 \\times 65 + 1 \\times 65 \\\\\n&= 6500 + 65 = \\mathbf{6565}\n\\end{aligned}\n\\]',
         '利用乘法分配律,拆分101为100+1,每一项都要乘65,不要漏乘。'),
        ('2.  \\( 99 \\times 42 \\)',
         '\\[\n\\begin{aligned}\n99 \\times 42 &= (100 - 1) \\times 42 \\\\\n&= 100 \\times 42 - 1 \\times 42 \\\\\n&= 4200 - 42 = \\mathbf{4158}\n\\end{aligned}\n\\]',
         '把99看成100-1,分配律展开,两项都要乘42。'),
        ('3.  \\( 25 \\times 36 \\)',
         '\\[\n\\begin{aligned}\n25 \\times 36 &= 25 \\times (4 \\times 9) \\\\\n&= (25 \\times 4) \\times 9 \\\\\n&= 100 \\times 9 = \\mathbf{900}\n\\end{aligned}\n\\]',
         '看到25想到拆出4,利用乘法结合律凑整。36=4×9。'),
        ('4.  \\( 125 \\times 24 \\times 25 \\)',
         '\\[\n\\begin{aligned}\n125 \\times 24 \\times 25 &= 125 \\times (8 \\times 3) \\times 25 \\\\\n&= (125 \\times 8) \\times (3 \\times 25) \\\\\n&= 1000 \\times 75 = \\mathbf{75000}\n\\end{aligned}\n\\]',
         '看到125和25,拆24为8×3,这样125×8=1000,凑整计算简便。也可以拆成 3×4×2,得到 (125×8)×(25×4)×3=1000×100×3=300000?不对,因为24=8×3,所以正确结果是75000。'),
        ('5.  \\( 25 \\times (40 + 8) \\)',
         '\\[\n\\begin{aligned}\n25 \\times (40 + 8) &= 25 \\times 40 + 25 \\times 8 \\\\\n&= 1000 + 200 = \\mathbf{1200}\n\\end{aligned}\n\\]',
         '乘法分配律,括号里每一项都要和25相乘,不要漏乘。'),
        ('6.  \\( 368 - 68 - 32 \\)',
         '\\[\n\\begin{aligned}\n368 - 68 - 32 &= 368 - (68 + 32) \\\\\n&= 368 - 100 = \\mathbf{268}\n\\end{aligned}\n\\]',
         '一个数连续减两个数,等于这个数减这两个数的和。添括号时,括号前面是减号,括号里面要变号。'),
    ]
    for q, ans, explain in ans_calc:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.size = Pt(12)
        p = doc.add_paragraph()
        run = p.add_run(ans)
        p = doc.add_paragraph()
        run = p.add_run(f'   {explain}')
        run.font.size = Pt(11)
        doc.add_paragraph()

    doc.add_paragraph()

    # 解决问题答案
    p = doc.add_heading('四、解决问题', level=1)
    p = doc.add_paragraph()
    run = p.add_run('解答:')
    p = doc.add_paragraph()
    run = p.add_run('1. 每件利润 = 售价 - 进价 = \\(\\mathbf{68 - 45 = 23}\\) 元')
    p = doc.add_paragraph()
    run = p.add_run('2. 总件数 = 总利润 \\(\\div\\) 每件利润 = \\(\\mathbf{920 \\div 23 = 40}\\) 件')
    p = doc.add_paragraph()
    run = p.add_run('**答案:这批衬衫一共有 \\(\\boxed{\\mathbf{40}}\\) 件**')
    p = doc.add_paragraph()
    run = p.add_run('   提示:题目中"星期一卖出12件"是干扰信息,本题不需要用到。做题时要注意区分有用信息和干扰信息。')
    run.font.size = Pt(11)

    doc.add_paragraph()

    # 评分标准
    p = doc.add_heading('评分标准', level=1)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '总分'
    hdr[1].text = '评价'
    hdr[2].text = '建议'
    row = table.rows[1].cells
    row[0].text = '90-100\n70-89\n<70'
    row[1].text = '🎉 优秀\n👍 良好\n📚 需要加强'
    row[2].text = '整理错题,考前复习\n针对错题知识点再练\n回归课本,重新理解概念'

    doc.add_paragraph()

    # 知识点复习提示
    p = doc.add_paragraph()
    run = p.add_run('📖 复习提示:做错的题目对应知识点,请回到错题本对照复习')
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run(f'生成日期:{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 保存
    output_path = '/home/wangfei/.openclaw/workspace/四年级数学错题专项复习卷第二套.docx'
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    output = create_math_test()
    print(f"Word文档已生成: {output}")