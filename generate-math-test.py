#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成多多四年级数学错题专项练习卷 - Word文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def create_math_test():
    doc = Document()
    
    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 标题
    title = doc.add_heading('四年级数学下册\n错题专项练习卷（第一套）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    
    # 信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('姓名：___________    用时：___________    得分：___________')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 第一部分：填空题
    heading1 = doc.add_heading('一、填空题（每空5分，共50分）', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('1. 用1、3、5、8四个数字和小数点组成：\n   - 最大的一位小数是 ____________\n   - 最小的两位小数是 ____________')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('2. 一个两位小数四舍五入后是6.8，这个两位小数：\n   - 最大是 ____________\n   - 最小是 ____________')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('3. 一个三位小数四舍五入后是9.50，这个三位小数：\n   - 最大是 ____________\n   - 最小是 ____________')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('4. 一块长方体木料，长8厘米，宽5厘米，高3厘米。截取一个面：\n   - 最大面的面积是 ____________ 平方厘米\n   - 最小面的面积是 ____________ 平方厘米')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('5. 火车每小时行驶v千米，行驶了s千米，需要 ____________ 小时。（用字母表示）')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 第二部分：选择题
    heading2 = doc.add_heading('二、选择题（每题5分，共20分）', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('6. 甲数是m，比乙数的3倍少n，乙数是（  ）')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('A. 3m - n    B. (m - n) ÷ 3    C. (m + n) ÷ 3    D. m ÷ 3 - n')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('你的选择：____')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('7. 学校买了a个足球，每个58元，买了b个篮球，每个85元，买足球比买篮球少花（  ）元')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('A. 85b - 58a    B. 58a - 85b    C. (85 - 58)(b - a)    D. 85b + 58a')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('你的选择：____')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('8. 小明今年a岁，爸爸比小明大28岁，10年后，爸爸比小明大（  ）岁')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('A. a + 28    B. 28    C. 38    D. a + 38')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('你的选择：____')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run('9. 长跑比赛，小刚用了m分钟，小强比小刚慢3分钟，小强用了（  ）分钟')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('A. m + 3    B. m - 3    C. 3m    D. m ÷ 3')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('你的选择：____')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 第三部分：解答题
    heading3 = doc.add_heading('三、解答题（每题10分，共30分）', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('10. 水果店每千克苹果a元，每千克香蕉b元。')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('   买4千克苹果比买3千克香蕉便宜多少元？用含有字母的式子表示。')
    run.font.size = Pt(12)
    p = doc.add_paragraph('\n\n\n')
    
    p = doc.add_paragraph()
    run = p.add_run('11. 一个长方形操场，长120米，宽80米。小明每分钟走v米，他绕操场走一圈需要多少分钟？用字母公式表示。如果v=50，计算需要多少分钟？')
    run.font.size = Pt(12)
    p = doc.add_paragraph('\n\n\n\n')
    
    p = doc.add_paragraph()
    run = p.add_run('12. 用0、2、3、5四个数字和小数点，写出符合要求的数（每个数字都要用上且只能用一次）：')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('   - 小于1且小数部分是三位的最大数是多少？')
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('   - 大于5且小数部分是三位的最小数是多少？')
    run.font.size = Pt(12)
    p = doc.add_paragraph('\n\n\n')
    
    # 分页
    doc.add_page_break()
    
    # 答案部分
    ans_title = doc.add_heading('参考答案与解析', 0)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 填空题答案
    p = doc.add_heading('一、填空题', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('1. ').bold = True
    run = p.add_run('最大的一位小数：')
    run = p.add_run('853.1').bold = True
    run = p.add_run('   最小的两位小数：')
    run = p.add_run('13.58').bold = True
    p = doc.add_paragraph('   解析：一位小数 → 小数点后只有1位，剩下3个数放整数部分，从大到小排列；'
                       '\n   两位小数 → 整数部分两位，最小不能以0开头，所以1开头，从小到大排列。')
    
    p = doc.add_paragraph()
    run = p.add_run('2. ').bold = True
    run = p.add_run('最大：')
    run = p.add_run('6.84').bold = True
    run = p.add_run('   最小：')
    run = p.add_run('6.75').bold = True
    p = doc.add_paragraph('   解析：四舍得到6.8 → 最大第二位是4；五入得到6.8 → 最小第二位是5。记住：最大四舍，最小五入。')
    
    p = doc.add_paragraph()
    run = p.add_run('3. ').bold = True
    run = p.add_run('最大：')
    run = p.add_run('9.504').bold = True
    run = p.add_run('   最小：')
    run = p.add_run('9.495').bold = True
    
    p = doc.add_paragraph()
    run = p.add_run('4. ').bold = True
    run = p.add_run('最大面：')
    run = p.add_run('40 平方厘米').bold = True
    run = p.add_run('   最小面：')
    run = p.add_run('15 平方厘米').bold = True
    p = doc.add_paragraph('   规律：最大面 = 最长两边相乘 (8×5=40)，最小面 = 最短两边相乘 (5×3=15)。')
    
    p = doc.add_paragraph()
    run = p.add_run('5. ').bold = True
    run = p.add_run('s ÷ v （或写成 $\\frac{s}{v}$）').bold = True
    p = doc.add_paragraph('   公式：时间 = 路程 ÷ 速度')
    
    doc.add_paragraph()
    
    # 选择题答案
    p = doc.add_heading('二、选择题', level=1)
    
    answers_choice = [
        ('6', 'C', '推导：甲数 = 乙数 × 3 - n → m = 3乙 - n → 3乙 = m + n → 乙 = (m + n) ÷ 3'),
        ('7', 'A', '"买足球比买篮球少花" = 篮球总价 - 足球总价 = 85b - 58a'),
        ('8', 'B', '年龄差永远不变，还是28岁'),
        ('9', 'A', '"慢" → 用时更多 → m + 3'),
    ]
    
    for num, ans, explain in answers_choice:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}. ').bold = True
        run = p.add_run(f'答案：{ans}').bold = True
        p = doc.add_paragraph(f'   {explain}')
    
    doc.add_paragraph()
    
    # 解答题答案
    p = doc.add_heading('三、解答题', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('10. ').bold = True
    run = p.add_run('解答：')
    p = doc.add_paragraph('   - 4千克苹果：4a 元')
    p = doc.add_paragraph('   - 3千克香蕉：3b 元')
    p = doc.add_paragraph('   - "便宜多少" = 香蕉总价 - 苹果总价 = ')
    run = p.add_run('   $\mathbf{3b - 4a}$ 元').bold = True
    p = doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('11. ').bold = True
    run = p.add_run('解答：')
    p = doc.add_paragraph('   - 操场周长 = (长 + 宽) × 2 = (120 + 80) × 2 = 400 米')
    p = doc.add_paragraph('   - 时间 = 路程 ÷ 速度 = $\mathbf{400 ÷ v}$ 分钟')
    p = doc.add_paragraph('   - 当 v = 50 时：400 ÷ 50 = ')
    run = p.add_run('   $\mathbf{8}$ 分钟').bold = True
    p = doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('12. ').bold = True
    run = p.add_run('解答：')
    p = doc.add_paragraph('   - 小于1，整数部分只能是0 → 剩下三个数字从大到小排 → ')
    run = p.add_run('   $\mathbf{0.532}$').bold = True
    p = doc.add_paragraph('   - 大于5，整数部分只能是5 → 剩下三个数字从小到大排 → ')
    run = p.add_run('   $\mathbf{5.023}$').bold = True
    
    doc.add_paragraph()
    
    # 评分表
    p = doc.add_heading('评分标准', level=1)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '总分'
    hdr_cells[1].text = '评价'
    hdr_cells[2].text = '建议'
    row_cells = table.rows[1].cells
    row_cells[0].text = '90-100\n70-89\n<70'
    row_cells[1].text = '🎉 优秀\n👍 良好\n📚 需要加强'
    row_cells[2].text = '掌握很好，考前看一遍错题\n再把错的对应知识点复习一遍\n回到错题本重新理解概念，再多练一套'
    
    doc.add_paragraph()
    
    # 备注
    p = doc.add_paragraph()
    run = p.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存
    output_path = '/home/wangfei/.openclaw/workspace/四年级数学错题专项练习卷第一套.docx'
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    output = create_math_test()
    print(f"Word文档已生成: {output}")
