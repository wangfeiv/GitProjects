#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
duoduo-study Word 复习卷生成模板

目标：
1. 生成纯文本 docx，避免乱码
2. 结构稳定：标题、信息栏、题目区、答案区
3. 每题带题型、难度、答案、解析
4. 适合作为后续 skill 直接复用的模板

使用示例：
python3 scripts/generate_review_docx.py data.json output.docx

JSON 输入结构示例：
{
  "title": "多多数学复习卷",
  "student_name": "多多",
  "subject": "数学",
  "grade": "小学四年级下学期",
  "note": "本卷根据错题库生成",
  "sections": [
    {
      "name": "填空题",
      "questions": [
        {
          "number": 1,
          "type": "填空题",
          "difficulty": "基础",
          "text": "一个三位小数四舍五入后是 8.60，这个三位小数最大是多少？",
          "answer": "8.604",
          "explanation": "最大是四舍得到，所以最后一位最大是4。"
        }
      ]
    }
  ]
}
"""

import json
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


SAFE_REPLACEMENTS = {
    "\\": "",
    "\\(" : "",
    "\\)" : "",
    "\\[" : "",
    "\\]" : "",
    "**": "",
    "__": "",
    "```": "",
    "$": "",
}


def sanitize_text(text):
    if text is None:
        return ""
    text = str(text)
    for old, new in SAFE_REPLACEMENTS.items():
        text = text.replace(old, new)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 去掉常见可能导致 Word 观感很差的技术性片段
    bad_fragments = [
        "<xml>", "</xml>", "{\\", "\\frac", "\\times", "\\div", "\\boxed", "\\begin", "\\end"
    ]
    for frag in bad_fragments:
        text = text.replace(frag, "")
    return text.strip()


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", size=12, bold=False, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(sanitize_text(text))
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def ensure_parent(path):
    parent = os.path.dirname(os.path.abspath(path))
    if parent:
        os.makedirs(parent, exist_ok=True)


def build_doc(data, output_path):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(12)

    title = sanitize_text(data.get("title", "多多复习卷"))
    student_name = sanitize_text(data.get("student_name", "多多"))
    subject = sanitize_text(data.get("subject", "未指定"))
    grade = sanitize_text(data.get("grade", "小学四年级"))
    note = sanitize_text(data.get("note", "本卷根据错题库生成"))
    sections = data.get("sections", [])

    add_paragraph(doc, title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"姓名：{student_name}    科目：{subject}    年级：{grade}    日期：{datetime.now().strftime('%Y-%m-%d')}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, note, size=11, color=(100, 100, 100))

    add_paragraph(doc, "")
    add_paragraph(doc, "试题部分", size=15, bold=True)

    for sec in sections:
        sec_name = sanitize_text(sec.get("name", "未命名题组"))
        add_paragraph(doc, sec_name, size=14, bold=True)
        questions = sec.get("questions", [])
        for q in questions:
            number = q.get("number", "")
            qtype = sanitize_text(q.get("type", "未分类"))
            difficulty = sanitize_text(q.get("difficulty", "未标注"))
            text = sanitize_text(q.get("text", ""))
            add_paragraph(doc, f"{number}. [{qtype}] [难度：{difficulty}] {text}")
            add_paragraph(doc, "")

    doc.add_page_break()
    add_paragraph(doc, "参考答案与解析", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for sec in sections:
        sec_name = sanitize_text(sec.get("name", "未命名题组"))
        add_paragraph(doc, sec_name, size=14, bold=True)
        questions = sec.get("questions", [])
        for q in questions:
            number = q.get("number", "")
            qtype = sanitize_text(q.get("type", "未分类"))
            difficulty = sanitize_text(q.get("difficulty", "未标注"))
            answer = sanitize_text(q.get("answer", "待补充"))
            explanation = sanitize_text(q.get("explanation", "待补充"))
            add_paragraph(doc, f"{number}. [{qtype}] [难度：{difficulty}] 答案：{answer}", bold=True)
            add_paragraph(doc, f"解析：{explanation}", size=11)
            add_paragraph(doc, "")

    add_paragraph(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", size=10, color=(120, 120, 120), align=WD_ALIGN_PARAGRAPH.RIGHT)

    ensure_parent(output_path)
    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 generate_review_docx.py <input.json> <output.docx>")
        sys.exit(1)

    input_json = sys.argv[1]
    output_docx = sys.argv[2]

    with open(input_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    build_doc(data, output_docx)
    print(output_docx)


if __name__ == "__main__":
    main()
